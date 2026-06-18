import streamlit as st
import pandas as pd
from datetime import datetime
import math
import requests
from requests.adapters import HTTPAdapter
from urllib3.util import Retry
import base64
from PIL import Image, ImageDraw
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import time

# --- 1. GLOBAL UI & ELEGANT ENTERPRISE CSS THEME ---
st.set_page_config(page_title="OptiStaff - Enterprise", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

    :root{
        --bg:#F1F5F9; --surface:#FFFFFF; --border:#E6E9F0; --border-strong:#D6DBE6;
        --text:#0F172A; --muted:#64748B;
        --primary:#4F46E5; --primary-hover:#4338CA; --primary-soft:#EEF0FF;
        --accent:#06B6D4;
        --success:#10B981; --warning:#F59E0B; --danger:#EF4444;
        --radius:14px; --shadow:0 1px 2px rgba(15,23,42,.04), 0 10px 26px rgba(15,23,42,.06);
    }

    html, body, [class*="css"]{
        font-family:'Inter',-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,Helvetica,Arial,sans-serif;
        -webkit-font-smoothing:antialiased;
    }

    .stApp{
        background:
            radial-gradient(1100px 520px at 100% -8%, #EEF2FF 0%, rgba(238,242,255,0) 55%),
            radial-gradient(900px 460px at -10% 0%, #ECFEFF 0%, rgba(236,254,255,0) 52%),
            var(--bg);
    }

    @keyframes fadeIn { 0% { opacity: 0; transform: translateY(10px); } 100% { opacity: 1; transform: translateY(0); } }
    .main .block-container {
        max-width: 1140px;
        padding-top: 2.2rem;
        padding-bottom: 4rem;
        animation: fadeIn 0.5s ease-out forwards;
    }

    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header[data-testid="stHeader"] { background: transparent !important; box-shadow: none !important; }
    [data-testid="stSidebarCollapsedControl"], [data-testid="collapsedControl"] {
        visibility: visible !important; opacity: 1 !important; z-index: 999999 !important;
    }
    [data-testid="stSidebarCollapsedControl"] svg, [data-testid="collapsedControl"] svg,
    [data-testid="stSidebarCollapsedControl"] button, [data-testid="collapsedControl"] button {
        color: #4F46E5 !important; fill: #4F46E5 !important;
    }

    /* Headings & teks utama: paksa gelap, lepas dari tema device (light/dark) */
    section[data-testid="stMain"], .main { color: #1E293B; }
    section[data-testid="stMain"] h1, .main h1 { color: #0F172A !important; font-weight: 700; letter-spacing: -0.02em; }
    section[data-testid="stMain"] h2, section[data-testid="stMain"] h3, section[data-testid="stMain"] h4,
    .main h2, .main h3, .main h4 { color: #0F172A !important; font-weight: 600; letter-spacing: -0.01em; }
    section[data-testid="stMain"] [data-testid="stMarkdownContainer"] p,
    .main [data-testid="stMarkdownContainer"] p { color: #1E293B; }
    [data-testid="stSidebar"] [data-testid="stCaptionContainer"] { color: #94A3B8 !important; }
    section[data-testid="stMain"] [data-testid="stCaptionContainer"] { color: var(--muted) !important; }

    /* Primary buttons */
    div.stButton > button, div.stDownloadButton > button, div.stFormSubmitButton > button {
        border-radius: 10px;
        border: 1px solid transparent;
        font-weight: 600;
        background-color: var(--primary) !important;
        color: white !important;
        transition: all 0.18s cubic-bezier(0.25, 0.1, 0.25, 1);
        padding: 0.6rem 1.1rem;
        box-shadow: 0 1px 2px rgba(79,70,229,0.25);
    }
    div.stButton > button:hover, div.stDownloadButton > button:hover, div.stFormSubmitButton > button:hover {
        background-color: var(--primary-hover) !important;
        transform: translateY(-1px);
        box-shadow: 0 8px 18px rgba(79, 70, 229, 0.28) !important;
    }
    div.stButton > button:active { transform: translateY(0) scale(0.99); }

    /* Secondary buttons */
    button[kind="secondary"] {
        background-color: var(--surface) !important;
        color: var(--text) !important;
        border: 1px solid var(--border-strong) !important;
        box-shadow: none !important;
    }
    button[kind="secondary"]:hover {
        background-color: #F8FAFC !important;
        border-color: var(--primary) !important;
        color: var(--primary) !important;
    }

    /* Inputs */
    .stTextInput input, .stTextArea textarea, .stNumberInput input,
    .stSelectbox div[data-baseweb="select"], .stMultiSelect div[data-baseweb="select"] {
        border-radius: 10px !important;
        border-color: var(--border-strong) !important;
        background-color: var(--surface) !important;
        color: var(--text) !important;
        caret-color: var(--primary) !important;
        transition: border-color 0.2s, box-shadow 0.2s;
    }
    .stTextInput input, .stTextArea textarea, .stNumberInput input {
        -webkit-text-fill-color: var(--text) !important;
    }
    .stTextInput input::placeholder, .stTextArea textarea::placeholder {
        color: #94A3B8 !important; -webkit-text-fill-color: #94A3B8 !important;
    }
    .stTextInput input:focus, .stTextArea textarea:focus, .stNumberInput input:focus {
        border-color: var(--primary) !important;
        box-shadow: 0 0 0 3px rgba(79,70,229,0.15) !important;
    }
    label { font-weight: 500 !important; color: #334155 !important; }

    /* Forms & expanders as cards */
    [data-testid="stForm"] {
        background: var(--surface);
        border: 1px solid var(--border);
        border-radius: var(--radius);
        padding: 1.5rem 1.5rem 0.6rem;
        box-shadow: var(--shadow);
    }
    [data-testid="stExpander"] {
        border: 1px solid var(--border) !important;
        border-radius: 12px !important;
        background: var(--surface);
        box-shadow: var(--shadow);
        overflow: hidden;
    }

    /* Tabs */
    .stTabs [data-baseweb="tab-list"] { gap: 0.35rem; border-bottom: 1px solid var(--border); }
    .stTabs [data-baseweb="tab"] {
        height: auto; padding: 0.55rem 1rem; border-radius: 10px 10px 0 0;
        color: var(--muted); font-weight: 600;
    }
    .stTabs [aria-selected="true"] { color: var(--primary) !important; background: var(--primary-soft); }
    .stTabs [data-baseweb="tab-highlight"] { background-color: var(--primary); }

    /* Radio as segmented control */
    div[role="radiogroup"] { gap: 0.4rem; }
    div[role="radiogroup"] label {
        background: var(--surface); border: 1px solid var(--border-strong); border-radius: 10px;
        padding: 0.35rem 0.85rem; transition: all 0.15s ease;
    }
    div[role="radiogroup"] label:hover { border-color: var(--primary); }

    /* Metrics */
    [data-testid="stMetric"] {
        background: var(--surface); border: 1px solid var(--border); border-radius: 12px;
        padding: 1rem 1.1rem; box-shadow: var(--shadow);
    }
    [data-testid="stMetricValue"] { color: var(--text); font-weight: 700; }

    /* Alerts & dataframe */
    [data-testid="stAlert"] { border-radius: 12px; }
    [data-testid="stDataFrame"] { border: 1px solid var(--border); border-radius: 12px; overflow: hidden; }

    hr { margin: 1.1rem 0; border-color: var(--border); }
    img { max-width: 100%; height: auto; border-radius: 12px; }

    /* Sidebar (elegant dark navy) */
    [data-testid="stSidebar"] { background: #0F172A; border-right: 1px solid #1E293B; }
    [data-testid="stSidebar"] * { color: #E2E8F0; }
    [data-testid="stSidebar"] .stButton > button {
        background-color: #1E293B !important; border: 1px solid #334155 !important;
        color: #E2E8F0 !important; box-shadow: none !important;
    }
    [data-testid="stSidebar"] .stButton > button:hover {
        background-color: #334155 !important; border-color: #475569 !important; color: #FFFFFF !important;
    }
    [data-testid="stSidebar"] [data-testid="stAlert"] {
        background: rgba(255,255,255,0.06); border: 1px solid rgba(255,255,255,0.10);
    }

    /* Custom brand blocks */
    .os-hero { text-align: center; margin: 0.5rem auto 1.6rem; }
    .os-logo { display: inline-flex; align-items: center; gap: 0.6rem; font-weight: 700; font-size: 1.7rem; color: var(--text); }
    .os-logo .dot { width: 34px; height: 34px; border-radius: 9px; background: linear-gradient(135deg, #4F46E5, #06B6D4); display: inline-block; box-shadow: 0 6px 16px rgba(79,70,229,.35); }
    .os-tag { color: var(--muted); margin-top: 0.5rem; font-size: 0.95rem; }
    .os-sb-brand { display: flex; align-items: center; gap: 0.55rem; font-size: 1.25rem; font-weight: 700; color: #FFFFFF !important; padding: 0.2rem 0 0.1rem; }
    .os-sb-brand .dot { width: 26px; height: 26px; border-radius: 7px; background: linear-gradient(135deg, #6366F1, #22D3EE); }

    /* Responsive: rapikan untuk layar HP */
    .stTabs [data-baseweb="tab-list"] { overflow-x: auto; }
    @media (max-width: 640px) {
        .main .block-container { max-width: 100% !important; padding-left: 1rem; padding-right: 1rem; padding-top: 1.2rem; }
        .os-hero { margin-top: 0.2rem; }
        .os-logo { font-size: 1.4rem; }
        .os-logo .dot { width: 28px; height: 28px; }
        h1 { font-size: 1.6rem !important; }
        h2, h3 { font-size: 1.15rem !important; }
        .stTabs [data-baseweb="tab"] { padding: 0.45rem 0.6rem; font-size: 0.9rem; }
        [data-testid="stForm"] { padding: 1.1rem 1.1rem 0.4rem; }
        [data-testid="stMetricValue"] { font-size: 1.4rem !important; }
    }
</style>
""", unsafe_allow_html=True)

# --- 2. SECURE MASTER STATE MEMORY INITIALIZATION ---
if 'broadcast_db' not in st.session_state: st.session_state.broadcast_db = {}
if 'logged_in' not in st.session_state: st.session_state.logged_in = False
if 'user_role' not in st.session_state: st.session_state.user_role = None
if 'user_name' not in st.session_state: st.session_state.user_name = None
if 'user_event_id' not in st.session_state: st.session_state.user_event_id = None
if 'reset_key' not in st.session_state: st.session_state.reset_key = 0

# --- 3. KONTROL JARINGAN & CACHE BUSTER API ---
http_session = requests.Session()
strategi_retry = Retry(total=4, backoff_factor=1, status_forcelist=[500, 502, 503, 504], raise_on_status=False)
http_session.mount("https://", HTTPAdapter(max_retries=strategi_retry))

API_URL = "https://script.google.com/macros/s/AKfycbyMyVBSKWB_zxJKUjX7shSvk9pQ0WxB7krb3JlH4JFodo-iIOhPeJeEEzN5aZILaLjh/exec"

def fetch_data(sheet_name):
    """
    Fungsi penarik data tangguh. Menggunakan timestamp agar GSheets tidak memberi data basi (Cache Buster).
    """
    try:
        url_bebas_cache = f"{API_URL}?sheet={sheet_name}&t={int(time.time() * 1000)}"
        res = http_session.get(url_bebas_cache, timeout=20)
        if res.status_code == 200:
            data = res.json()
            if data: return pd.DataFrame(data)
        return pd.DataFrame()
    except:
        return pd.DataFrame()

# ==========================================
# 4. FUNGSI PENDUKUNG LINTAS MODUL
# ==========================================
def hitung_jarak(lat1, lon1, lat2, lon2):
    R = 6371000
    phi1, phi2 = math.radians(lat1), math.radians(lat2)
    dphi, dlam = math.radians(lat2-lat1), math.radians(lon2-lon1)
    a = math.sin(dphi/2)**2 + math.cos(phi1)*math.cos(phi2) * math.sin(dlam/2)**2
    return R * (2 * math.atan2(math.sqrt(a), math.sqrt(1-a)))

def terjemahkan_koordinat_ke_jalan(event_id, nama_event_gsheet):
    ev_clean = f"{event_id}".upper().strip()
    nm_clean = f"{nama_event_gsheet}".upper().strip()
    if "EVT01" in ev_clean or "MUSIC" in nm_clean or "FEST" in nm_clean:
        return "Mall Citraland Grogol, Jl. Letjen S. Parman, Tanjung Duren, Jakarta Barat"
    elif "EVT02" in ev_clean or "PSSI" in nm_clean or "GBK" in nm_clean:
        return "Stadion Utama Gelora Bung Karno (GBK), Pintu 1 Senayan, Jakarta Pusat"
    elif "EVT03" in ev_clean or "PRJ" in nm_clean:
        return "JIExpo Kemayoran, Gedung Pusat Niaga lt. 1, Kemayoran, Jakarta Pusat"
    return "Zonasi Kompleks Operasional Event"

def bersihkan_koordinat_radar(val, is_latitude=True):
    try:
        s = f"{val}".replace(',', '').replace('.', '').strip()
        if not s: return -6.2181 if is_latitude else 106.8024
        is_negative = s.startswith('-')
        if is_negative: s = s[1:]
        num = float(s)
        if is_latitude:
            while num > 9.0: num /= 10.0
            return -num if is_negative else num
        else:
            while num > 180.0: num /= 10.0
            if num < 90.0:
                while num < 100.0: num *= 10.0
            return -num if is_negative else num
    except: return -6.2181 if is_latitude else 106.8024

def buat_word_lpj_apple_style(df, event_id, info_event, df_tasks):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_m = p_meta.add_run("OPTI STAFF CLOUD SYSTEM\nExecutive Report Management")
    run_m.font.size = Pt(8)
    run_m.font.color.rgb = RGBColor(140, 140, 142)

    p_title = doc.add_paragraph()
    run_title = p_title.add_run("Laporan Pertanggungjawaban")
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(29, 29, 31)

    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(f"Project ID: {event_id} — {info_event.get('nama_event', 'Exhibition Pro Master')}")
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(100, 100, 100)

    p_line = doc.add_paragraph()
    p_line.add_run("—" * 65).font.color.rgb = RGBColor(210, 210, 215)

    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("1. Ringkasan Operasional Lapangan")
    run_h1.font.size = Pt(14)
    run_h1.font.bold = True

    p_summary = doc.add_paragraph()
    total_log = len(df)
    total_valid = len(df[df['status_geotag'].astype('str').str.contains("MATCH", na=False)])

    raw_loc = info_event.get('lokasi_target', '-6.2181, 106.8024')
    try: lt_p, ln_p = map(float, raw_loc.split(","))
    except: lt_p, ln_p = -6.2181, 106.8024

    c_lat_clean = bersihkan_koordinat_radar(lt_p, is_latitude=True)
    c_lon_clean = bersihkan_koordinat_radar(ln_p, is_latitude=False)
    alamat_nyata_jalan = terjemahkan_koordinat_ke_jalan(event_id, info_event.get('nama_event',''))

    p_summary.add_run(f"Waktu Penyerahan Dokumen: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    p_summary.add_run(f"Pusat Koordinat Geofence: {c_lat_clean:.6f}, {c_lon_clean:.6f} ({alamat_nyata_jalan})\n")
    p_summary.add_run(f"Total Aktivitas Personel Lapangan: {total_log} Transaksi\n")
    p_summary.add_run(f"Hasil Validasi Radar Satelit: {total_valid} Valid\n")

    h_tasks = doc.add_heading(level=1)
    h_tasks.add_run("2. Status Capaian Tugas Operasional Divisi")
    p_tasks = doc.add_paragraph()

    if df_tasks is None or df_tasks.empty:
        p_tasks.add_run("Tidak ada penugasan to-do list khusus yang tercatat pada proyek ini.")
    else:
        df_event_tasks = df_tasks[df_tasks['event_id'] == event_id]
        t_total = len(df_event_tasks)
        t_done = len(df_event_tasks[df_event_tasks['status'] == "Disetujui"]) # Diubah menjadi Disetujui
        p_tasks.add_run(f"Total Instansi Tugas Terdistribusi: {t_total} | Berhasil Diselesaikan & Disetujui TL: {t_done} | Belum Selesai/Belum Disetujui: {t_total - t_done}\n")
        if t_done < t_total:
            r_warn = p_tasks.add_run("CATATAN AUDIT: Terdapat tugas divisi yang tidak diselesaikan oleh kru lapangan atau belum disetujui TL.")
            r_warn.font.color.rgb = RGBColor(255, 59, 48)
            r_warn.font.bold = True

    h3 = doc.add_heading(level=1)
    h3.add_run("3. Log Detail Presensi Masuk & Pulang").font.size = Pt(14)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    titles = ["Nama Personel", "Waktu Kirim", "Tipe Shift", "Keamanan Geotag", "Laporan Progress Kerja"]
    for idx, text in enumerate(titles):
        hdr_cells[idx].text = text
        hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        bg_shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F7"/>')
        hdr_cells[idx]._tc.get_or_add_tcPr().append(bg_shading)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = f"{row.get('crew_name', '-')}"
        row_cells[1].text = f"{row.get('timestamp', '-')}"
        row_cells[2].text = f"{row.get('tipe_absen', '-')}"
        status_gps = f"{row.get('status_geotag', '-')}"
        row_cells[3].text = status_gps
        if "OUT OF RANGE" in status_gps:
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 59, 48)
            row_cells[3].paragraphs[0].runs[0].font.bold = True
        row_cells[4].text = f"{row.get('status_pekerjaan', '-')}"

    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io.getvalue()

# ==========================================
# 5. AUTHENTICATION GATEWAY
# ==========================================
if not st.session_state.logged_in:
    st.markdown("""
    <div class="os-hero">
        <div class="os-logo"><span class="dot"></span> OptiStaff</div>
        <div class="os-tag">Sistem Manajemen Kehadiran Terintegrasi &middot; Geofencing &amp; Database Divisi</div>
    </div>
    """, unsafe_allow_html=True)

    _lc1, _lcmid, _lc2 = st.columns([1, 1.6, 1])
    with _lcmid:
        with st.form("login_gate"):
            input_email = st.text_input("Email Pengguna")
            input_pass = st.text_input("Kata Sandi", type="password")
            if st.form_submit_button("Masuk ke Sistem", use_container_width=True):
                with st.spinner("Memvalidasi kredensial dari Cloud..."):
                    if input_email == "eo@optistaff.com" and input_pass == "adminpro":
                        st.session_state.update({"logged_in": True, "user_role": "EO (Organizer)", "user_name": "Administrator Pusat", "user_event_id": "ALL"})
                        time.sleep(0.5); st.rerun()
                    else:
                        df_crews_db = fetch_data("crews")
                        if not df_crews_db.empty and input_email in df_crews_db['email'].values:
                            u_info = df_crews_db[df_crews_db['email'] == input_email].iloc[0]
                            if f"{u_info.get('password', '')}" == input_pass:
                                if f"{u_info.get('status', 'Aktif')}" == 'Aktif':
                                    st.session_state.update({"logged_in": True, "user_role": u_info['role'], "user_name": u_info['nama'], "user_event_id": u_info['event_id']})
                                    time.sleep(0.5); st.rerun()
                                else: st.error("Akses ditolak. Masa aktif akun Anda telah berakhir.")
                            else: st.error("Kredensial tidak valid.")
                        else: st.error("Email tidak ditemukan atau server GSheets sibuk.")
    st.stop()

# --- SIDEBAR CONTROL PANEL ---
st.sidebar.markdown('<div class="os-sb-brand"><span class="dot"></span> OptiStaff</div>', unsafe_allow_html=True)
st.sidebar.caption("Operations Control System")
st.sidebar.markdown("---")
st.sidebar.markdown(f"**Pengguna:** {st.session_state.user_name}")
st.sidebar.info(f"Hak Akses: {st.session_state.user_role}")
if st.session_state.user_event_id and st.session_state.user_event_id != "ALL":
    st.sidebar.success(f"ID Proyek: {st.session_state.user_event_id}")
else: st.sidebar.warning("Mode: Akses Administrator")

if st.sidebar.button("Keluar Sistem", type="secondary", use_container_width=True):
    with st.spinner("Mengakhiri sesi..."):
        time.sleep(0.5)
        st.session_state.logged_in = False
        st.rerun()

# ==========================================
# 6. INTERFACE ROLE: EO (ORGANIZER)
# ==========================================
if st.session_state.user_role == "EO (Organizer)":
    st.title("Portal Manajemen Pusat")
    tab_ev, tab_cr, tab_mon = st.tabs(["Manajemen Proyek", "Registrasi Personel", "Penutupan Proyek"])

    with tab_ev:
        st.subheader("Pendaftaran Proyek & Pembuatan Divisi Operasional")
        df_events = fetch_data("events")
        if df_events.empty: df_events = pd.DataFrame(columns=["event_id", "nama_event", "lokasi_target", "radius_meter", "status", "status_lpj", "divisions"])

        _ev_total = len(df_events)
        _ev_aktif = len(df_events[df_events['status'] == 'Aktif']) if 'status' in df_events.columns else 0
        _ec1, _ec2 = st.columns(2)
        _ec1.metric("Total Proyek", _ev_total)
        _ec2.metric("Proyek Aktif", _ev_aktif)

        with st.form("add_event_form", clear_on_submit=True):
            ev_id = st.text_input("ID Proyek:", value=f"EVT0{len(df_events)+1}")
            ev_nama = st.text_input("Nama Proyek:")
            ev_loc = st.text_input("Koordinat Pusat (Lat, Lon):", value="-6.2181, 106.8024")
            ev_rad = st.slider("Radius Geofence (Meter):", 50, 500, 100)

            st.markdown("**Konfigurasi Divisi Eksekusi Lapangan**")
            divisi_default = ["Ticketing", "Logistik", "Keamanan", "Konsumsi", "Dokumentasi", "Stage Management"]
            div_terpilih = st.multiselect("Struktur Divisi Standar Sistem:", divisi_default, default=divisi_default)
            div_kustom = st.text_input("Divisi Kustom Tambahan (Pisahkan dengan koma):", value="")

            if st.form_submit_button("Aktifkan Proyek & Divisi", use_container_width=True) and ev_nama:
                with st.spinner("Memproses transmisi data ke Cloud..."):
                    list_div_final = [d.strip() for d in div_terpilih]
                    if div_kustom:
                        list_div_final.extend([d.strip() for d in div_kustom.split(",") if d.strip()])

                    payload = {
                        "target_sheet": "events",
                        "event_id": ev_id,
                        "nama_event": ev_nama,
                        "lokasi_target": ev_loc,
                        "radius_meter": ev_rad,
                        "status": "Aktif",
                        "status_lpj": "Belum Diserahkan",
                        "divisions": ",".join(list_div_final)
                    }
                    http_session.post(API_URL, json=payload)
                    st.success(f"Proyek '{ev_nama}' berhasil diinisiasi pada Database Cloud.")
                    time.sleep(1); st.rerun()
        st.markdown("**Daftar Proyek Aktif**")
        st.dataframe(df_events, use_container_width=True)

    with tab_cr:
        st.subheader("Distribusi Personel & Pemetaan Divisi")
        df_crews = fetch_data("crews")
        if df_crews.empty: df_crews = pd.DataFrame(columns=["crew_id", "nama", "email", "role", "event_id", "password", "status", "division"])

        df_ev_list = fetch_data("events")
        if not df_ev_list.empty and 'event_id' in df_ev_list.columns:
            list_event_id = df_ev_list["event_id"].tolist()
        else: list_event_id = ["EVT01"]

        st.markdown("**1. Pilih Proyek Penugasan Terlebih Dahulu:**")
        c_ev_target = st.selectbox("Target Proyek:", list_event_id, label_visibility="collapsed")

        # PENGAMAN GANDA: Mencocokkan data GSheet meskipun kolomnya dinamai 'divisions' (typo) atau 'division'
        try:
            selected_ev_info = df_ev_list[df_ev_list['event_id'] == c_ev_target].iloc[0]
            col_target_div = 'divisions' if 'divisions' in selected_ev_info else ('division' if 'division' in selected_ev_info else None)
            if col_target_div:
                raw_div = selected_ev_info.get(col_target_div, '')
                if pd.isna(raw_div) or str(raw_div).strip() == "" or str(raw_div).lower() == "nan":
                    div_options = ["Ticketing", "Logistik", "Keamanan", "Konsumsi", "Dokumentasi", "Stage Management"]
                else:
                    div_options = [d.strip() for d in str(raw_div).split(',') if d.strip()]
            else:
                div_options = ["Ticketing", "Logistik", "Keamanan", "Konsumsi", "Dokumentasi", "Stage Management"]
        except:
            div_options = ["Ticketing", "Logistik", "Keamanan", "Konsumsi", "Dokumentasi", "Stage Management"]

        st.markdown("**2. Konfigurasi Hak Akses:**")
        c_role = st.radio("Tingkat Akses Personel:", ["Team Leader", "Crew"], horizontal=True)

        with st.form("add_crew_form", clear_on_submit=True):
            c_id = st.text_input("ID Personel:", value=f"CRW0{len(df_crews)+1}")
            c_nama = st.text_input("Nama Lengkap:")
            c_email = st.text_input("Email Login:")
            c_pass = st.text_input("Kata Sandi:", value="12345")

            if c_role == "Crew":
                c_div = st.selectbox("Alokasi Divisi Operasional:", div_options)
            else:
                st.info("Otoritas Team Leader otomatis mencakup seluruh divisi pada proyek ini.")
                c_div = "Team Leader"

            if st.form_submit_button("Daftarkan Personel", use_container_width=True) and c_nama and c_email:
                with st.spinner("Mentransmisikan data kredensial ke Cloud..."):
                    final_div = "Team Leader" if c_role == "Team Leader" else c_div
                    payload = {"target_sheet": "crews", "crew_id": c_id, "nama": c_nama, "email": c_email, "role": c_role, "event_id": c_ev_target, "password": c_pass, "status": "Aktif", "division": final_div}
                    http_session.post(API_URL, json=payload)
                    st.success(f"Personel {c_nama} terdaftar di Cloud database sebagai {c_role} pada divisi: {final_div}.")
                    time.sleep(1.5); st.rerun()

        st.markdown("**Direktori Personel**")
        st.dataframe(df_crews, use_container_width=True)

    with tab_mon:
        st.subheader("Penyelesaian & Terminasi Proyek")
        df_all_att = fetch_data("attendances")
        df_events_check = fetch_data("events")

        if not df_events_check.empty and 'status' in df_events_check.columns:
            unique_events_aktif = df_events_check[df_events_check['status'] == 'Aktif']["event_id"].tolist()

            if unique_events_aktif:
                c_close1, c_close2 = st.columns([3, 1])
                with c_close1:
                    event_target_close = st.selectbox("Pilih Proyek untuk Diterminasi:", unique_events_aktif, key="close_ev")
                with c_close2:
                    st.markdown("<div style='margin-top: 28px;' class='hidden-mobile'></div>", unsafe_allow_html=True)
                    tombol_close = st.button("Tutup Proyek", type="primary", use_container_width=True)

                info_event = df_events_check[df_events_check['event_id'] == event_target_close].iloc[0]
                status_lpj_saat_ini = info_event.get('status_lpj', 'Belum Diserahkan')
                st.info(f"Status Dokumen LPJ: {status_lpj_saat_ini}")

                if tombol_close:
                    if status_lpj_saat_ini != "Sudah Diserahkan":
                        st.error("Terminasi ditolak. Menunggu pengesahan laporan dari pengawas lapangan (Team Leader).")
                    else:
                        with st.spinner("Memproses pengarsipan..."):
                            res_close = http_session.post(API_URL, json={"action": "complete_event", "event_id": event_target_close})
                            if res_close.json().get("status") == "success":
                                st.success(f"Proyek {event_target_close} resmi diterminasi dan diarsipkan."); time.sleep(1); st.rerun()

                st.markdown("---")
                st.subheader("Riwayat Log Kehadiran Proyek")
                if not df_all_att.empty:
                    selected_event_filter = st.selectbox("Filter berdasarkan Proyek:", df_all_att['event_id'].unique().tolist())
                    df_filtered = df_all_att[df_all_att['event_id'] == selected_event_filter]
                    if not df_filtered.empty:
                        df_lpj_eo = df_filtered[['attendance_id', 'crew_name', 'tipe_absen', 'timestamp', 'status_geotag', 'status_pekerjaan']].copy()
                        def color_eo_status(val):
                            if "OUT OF RANGE" in f"{val}": return 'background-color: #fde8e8; color: #9b1c1c;'
                            return 'background-color: #eafbf0; color: #155724;'
                        _eo_styler = df_lpj_eo.style
                        _eo_apply = _eo_styler.map if hasattr(_eo_styler, "map") else _eo_styler.applymap
                        st.dataframe(_eo_apply(color_eo_status, subset=['status_geotag']), use_container_width=True)

# ==========================================
# 7. INTERFACE ROLE: TEAM LEADER
# ==========================================
elif st.session_state.user_role == "Team Leader":
    active_ev_id = st.session_state.user_event_id
    st.title("Dashboard Pengawas Operasional")
    st.caption(f"ID Proyek Aktif: {active_ev_id}")

    tab_tl_mon, tab_tl_divisi, tab_tl_add = st.tabs(["Pemantauan Radar", "Manajemen Tugas Divisi", "Manajemen Kru Lokal"])

    df_tasks_cloud = fetch_data("tasks")

    with tab_tl_mon:
        df_attendance = fetch_data("attendances")
        if not df_attendance.empty and 'event_id' in df_attendance.columns:
            df_attendance = df_attendance[df_attendance['event_id'] == active_ev_id]

        df_ev_all = fetch_data("events")
        if not df_ev_all.empty:
            info_ev_df = df_ev_all[df_ev_all["event_id"] == active_ev_id]
            info_ev = info_ev_df.iloc[0] if not info_ev_df.empty else {"nama_event": "Exhibition"}
        else: info_ev = {"nama_event": "Exhibition"}

        st.subheader("Log Kehadiran Personel")
        if df_attendance.empty: st.info("Belum ada data kehadiran terdeteksi pada proyek ini.")
        else:
            df_display = df_attendance.copy()
            _m_total = len(df_display)
            _m_valid = len(df_display[df_display['status_geotag'].astype('str').str.contains("MATCH", na=False)])
            _m_out = _m_total - _m_valid
            _mc1, _mc2, _mc3 = st.columns(3)
            _mc1.metric("Total Kehadiran", _m_total)
            _mc2.metric("Lokasi Valid", _m_valid)
            _mc3.metric("Di Luar Radar", _m_out)
            st.dataframe(df_display[[c for c in df_display.columns if c != 'foto_selfie_url']], use_container_width=True)

            st.markdown("---")
            st.subheader("Radar Geolocation")
            df_map = df_display[['latitude', 'longitude', 'crew_name', 'status_geotag']].copy()
            df_map['latitude'] = df_map['latitude'].apply(lambda x: bersihkan_koordinat_radar(x, is_latitude=True))
            df_map['longitude'] = df_map['longitude'].apply(lambda x: bersihkan_koordinat_radar(x, is_latitude=False))
            df_map = df_map.dropna(subset=['latitude', 'longitude'])
            if not df_map.empty: st.map(df_map, latitude='latitude', longitude='longitude', size=40)

            st.markdown("---")
            st.subheader("Otorisasi Laporan Akhir (LPJ)")
            apple_docx = buat_word_lpj_apple_style(df_display, active_ev_id, info_ev, df_tasks_cloud)

            c_doc1, c_doc2 = st.columns(2)
            with c_doc1:
                st.download_button(label="Unduh Dokumen Draf (.docx)", data=apple_docx, file_name=f"LPJ_{active_ev_id}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="secondary", use_container_width=True)
            with c_doc2:
                if st.button("Sahkan & Kirim Data Audit ke Pusat", type="primary", use_container_width=True):
                    with st.spinner("Mengirim persetujuan..."):
                        res_lpj = http_session.post(API_URL, json={"action": "submit_lpj", "event_id": active_ev_id}).json()
                        if res_lpj.get("status") == "success": st.success("Dokumen LPJ berhasil diotorisasi.")

            st.markdown("---")
            st.subheader("Verifikasi Visual Kamera")
            df_display['pilihan_id'] = df_display['crew_name'] + " (" + df_display['attendance_id'] + ")"
            id_terpilih = st.selectbox("Pilih Personel:", df_display['pilihan_id'].tolist())
            row_data = df_display[df_display['pilihan_id'] == id_terpilih].iloc[0]
            if row_data['foto_selfie_url'] and "base64," in f"{row_data['foto_selfie_url']}":
                c_img, c_txt = st.columns([1, 2])
                with c_img: st.image(row_data['foto_selfie_url'], use_container_width=True)
                with c_txt:
                    status_geo = f"{row_data.get('status_geotag', '')}"
                    if "OUT OF RANGE" in status_geo: st.error(f"PELANGGARAN LOKASI: {status_geo}")
                    elif "MATCH" in status_geo: st.success(f"LOKASI VALID: {status_geo}")
                    else: st.info(f"Status Geofence: {status_geo}")
                    st.info(f"Tipe Shift: {row_data.get('tipe_absen','-')}")
                    st.write(f"**Laporan Pekerjaan:** {row_data.get('status_pekerjaan','-')}")

    with tab_tl_divisi:
        st.subheader("Pusat Distribusi Tugas Operasional Cloud")
        df_crews_all = fetch_data("crews")
        df_ev_all_raw = fetch_data("events")

        try:
            selected_ev_row = df_ev_all_raw[df_ev_all_raw['event_id'] == active_ev_id].iloc[0]
            col_target_div = 'divisions' if 'divisions' in selected_ev_row else ('division' if 'division' in selected_ev_row else None)
            if col_target_div:
                raw_div_tl = selected_ev_row.get(col_target_div, '')
                if pd.isna(raw_div_tl) or str(raw_div_tl).strip() == "" or str(raw_div_tl).lower() == "nan":
                    div_tersedia = ["Ticketing", "Logistik", "Keamanan", "Konsumsi", "Dokumentasi", "Stage Management"]
                else:
                    div_tersedia = [d.strip() for d in str(raw_div_tl).split(',') if d.strip()]
            else:
                div_tersedia = ["Ticketing", "Logistik", "Keamanan", "Konsumsi", "Dokumentasi", "Stage Management"]
        except:
            div_tersedia = ["Ticketing", "Logistik", "Keamanan", "Konsumsi", "Dokumentasi", "Stage Management"]

        target_div = st.selectbox("Pilih Divisi Operasional:", div_tersedia)

        # PENGAMAN GANDA: Auto-koreksi nama kolom 'division' atau 'divisions' di sheet crews
        col_crew_div = 'division' if 'division' in df_crews_all.columns else ('divisions' if 'divisions' in df_crews_all.columns else None)
        if not df_crews_all.empty and col_crew_div:
            kru_divisi_ini = df_crews_all[(df_crews_all[col_crew_div] == target_div) & (df_crews_all['event_id'] == active_ev_id) & (df_crews_all['role'] == 'Crew')]["nama"].tolist()
        else: kru_divisi_ini = []

        st.markdown("---")
        st.subheader("1. Distribusi Tugas Baru")
        if not kru_divisi_ini:
            st.warning("Sistem tidak mendeteksi kru murni pada divisi ini.")
        else:
            st.text(", ".join(kru_divisi_ini))

            with st.form("form_distribusi_task", clear_on_submit=True):
                raw_tasks_input = st.text_area("Parameter Tugas (1 baris mewakili 1 instruksi kerja):", value="Verifikasi alat scanner barcode")
                ref_cam = st.camera_input("Foto Referensi Arahan untuk Crew (Opsional)")

                if st.form_submit_button("Alokasikan Tugas", use_container_width=True):
                    list_tugas_kasar = [t.strip() for t in raw_tasks_input.split("\n") if t.strip()]
                    if list_tugas_kasar:
                        with st.spinner("Mengunggah tugas..."):
                            ref_b64 = ""
                            if ref_cam:
                                raw_img = Image.open(io.BytesIO(ref_cam.getvalue()))
                                raw_img.thumbnail((300, 300))
                                buffered_io = io.BytesIO()
                                raw_img.save(buffered_io, format="JPEG", quality=50)
                                ref_b64 = f"data:image/jpeg;base64,{base64.b64encode(buffered_io.getvalue()).decode('utf-8')}"

                            for idx, task_name in enumerate(list_tugas_kasar):
                                crew_target_name = kru_divisi_ini[idx % len(kru_divisi_ini)]
                                payload_task = {
                                    "target_sheet": "tasks", "task_id": f"TSK-{datetime.now().strftime('%H%M%S')}-{idx}",
                                    "event_id": active_ev_id, "division": target_div, "crew_name": crew_target_name,
                                    "task_name": task_name, "status": "Belum Selesai", "photo_report": "", "reference_photo": ref_b64, "tl_comment": ""
                                }
                                http_session.post(API_URL, json=payload_task)
                            st.success("Tugas dialokasikan."); time.sleep(1); st.rerun()

        st.markdown("---")
        st.subheader("2. Verifikasi Laporan Tugas Kru")
        if not df_tasks_cloud.empty and 'status' in df_tasks_cloud.columns:
            pending_tasks = df_tasks_cloud[(df_tasks_cloud['event_id'] == active_ev_id) & (df_tasks_cloud['division'] == target_div) & (df_tasks_cloud['status'] == "Menunggu Verifikasi")]
            if not pending_tasks.empty:
                for idx, row in pending_tasks.iterrows():
                    with st.expander(f"Kru: {row['crew_name']} ➡️ Tugas: {row['task_name']}", expanded=True):
                        if pd.notna(row.get('photo_report', '')) and "base64" in str(row['photo_report']):
                            try:
                                b64_str = str(row['photo_report']).split("base64,")[1]
                                st.image(Image.open(io.BytesIO(base64.b64decode(b64_str))), caption="📸 Bukti Selesai Kerja dari Kru")
                            except: pass

                        komentar_input = st.text_input("Komentar Evaluasi TL (Wajib diisi jika akan di Reject):", key=f"komentar_{row['task_id']}")
                        c1, c2 = st.columns(2)
                        with c1:
                            if st.button("✅ APPROVE (Selesai)", key=f"app_{row['task_id']}", use_container_width=True):
                                with st.spinner("Mensahkan tugas..."):
                                    payload = {"action": "update_task", "task_id": row['task_id'], "status": "Disetujui", "tl_comment": komentar_input}
                                    http_session.post(API_URL, json=payload)
                                    st.rerun()
                        with c2:
                            if st.button("❌ REJECT (Minta Ulang)", key=f"rej_{row['task_id']}", use_container_width=True):
                                if not komentar_input: st.error("Harap berikan komentar alasan reject agar kru bisa memperbaikinya.")
                                else:
                                    with st.spinner("Menolak tugas..."):
                                        payload = {"action": "update_task", "task_id": row['task_id'], "status": "Ditolak", "tl_comment": komentar_input}
                                        http_session.post(API_URL, json=payload)
                                        st.rerun()
            else: st.info("Tidak ada laporan baru yang menunggu verifikasi Anda.")

            st.markdown("---")
            st.subheader("3. Status Seluruh Pekerjaan Divisi")
            all_div_tasks = df_tasks_cloud[(df_tasks_cloud['event_id'] == active_ev_id) & (df_tasks_cloud['division'] == target_div)]
            st.dataframe(all_div_tasks[['crew_name', 'task_name', 'status', 'tl_comment']], use_container_width=True)

    with tab_tl_add:
        st.subheader("Rekrutmen Personel Lokal")
        with st.form("form_add_by_tl", clear_on_submit=True):
            n_id = st.text_input("ID Personel:", value=f"CRW-{datetime.now().strftime('%H%M')}")
            n_nama = st.text_input("Nama Lengkap:")
            n_email = st.text_input("Email Login:")
            n_pass = st.text_input("Kata Sandi (Password):", value="12345")
            n_div = st.selectbox("Alokasi Divisi Operasional:", div_tersedia)
            if st.form_submit_button("Daftarkan Personel", type="primary", use_container_width=True) and n_nama and n_email:
                with st.spinner("Mengirim data..."):
                    payload_tl = {"target_sheet": "crews", "crew_id": n_id, "nama": n_nama, "email": n_email, "role": "Crew", "event_id": active_ev_id, "password": n_pass, "status": "Aktif", "division": n_div}
                    http_session.post(API_URL, json=payload_tl)
                    st.success("Personel berhasil didaftarkan langsung ke database."); time.sleep(1); st.rerun()

# ==========================================
# 8. INTERFACE ROLE: CREW (PERSONEL LAPANGAN)
# ==========================================
else:
    active_ev_id = st.session_state.user_event_id
    st.title("Portal Operasional Lapangan")

    df_crews_all = fetch_data("crews")
    try:
        my_row_info = df_crews_all[df_crews_all['nama'] == st.session_state.user_name].iloc[0]
        col_crew_div = 'division' if 'division' in my_row_info else ('divisions' if 'divisions' in my_row_info else 'Umum')
        my_division = my_row_info.get(col_crew_div, 'Umum')
    except: my_division = "Umum"

    st.caption(f"Personel: {st.session_state.user_name} | Alokasi Divisi: {my_division}")

    st.markdown("---")
    menu_kru = st.radio("Pilih Modul Lapangan:", ["Absen Masuk", "Absen Pulang", "To-Do List Divisi"], horizontal=True)

    if menu_kru in ["Absen Masuk", "Absen Pulang"]:
        st.subheader(f"📡 Validasi Presensi Geofence ({menu_kru})")
        df_event_kru = fetch_data("events")
        try:
            df_event_kru = df_event_kru[df_event_kru['event_id'] == active_ev_id]
            lokasi_target = df_event_kru.iloc[0]['lokasi_target']
            rad_t = float(df_event_kru.iloc[0]['radius_meter'])
        except:
            lokasi_target = "-6.2181, 106.8024"
            rad_t = 100

        with st.container():
            foto_file = st.camera_input("Otorisasi Wajah")
            c1, c2 = st.columns(2)
            with c1: mock_lat = st.number_input("Latitude:", value=-6.2181, format="%.5f")
            with c2: mock_lon = st.number_input("Longitude:", value=106.8024, format="%.5f")
            laporan_kerja = st.text_area("Laporan Operasional Shift:")

            if st.button(f"Kirim Data {menu_kru}", type="primary", use_container_width=True) and foto_file:
                with st.spinner("Memproses transmisi geolokasi..."):
                    waktu_sekarang = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    img = Image.open(io.BytesIO(foto_file.getvalue()))
                    canvas = ImageDraw.Draw(img)
                    canvas.rectangle(((0, 0), (img.size[0], 35)), fill="black")
                    canvas.text((10, 10), f"AUTH: {waktu_sekarang} | {menu_kru.upper()} | PRJ: {active_ev_id}", fill="#00ffcc")

                    img.thumbnail((300, 300))
                    buffered_io = io.BytesIO()
                    img.save(buffered_io, format="JPEG", quality=50)
                    data_url_foto = f"data:image/jpeg;base64,{base64.b64encode(buffered_io.getvalue()).decode('utf-8')}"

                    try: lat_t, lon_t = map(float, str(lokasi_target).split(","))
                    except: lat_t, lon_t = -6.2181, 106.8024

                    c_lat_final, c_lon_final = bersihkan_koordinat_radar(mock_lat, True), bersihkan_koordinat_radar(mock_lon, False)
                    jarak = hitung_jarak(c_lat_final, c_lon_final, lat_t, lon_t)

                    is_valid = jarak <= rad_t
                    status_absen = "MATCH (Valid)" if is_valid else "OUT OF RANGE (Fraud Suspect)"

                    payload = {"target_sheet": "attendances", "attendance_id": f"ATT-{datetime.now().strftime('%H%M%S')}", "event_id": active_ev_id, "crew_name": st.session_state.user_name, "timestamp": waktu_sekarang, "latitude": c_lat_final, "longitude": c_lon_final, "status_geotag": status_absen, "status_pekerjaan": laporan_kerja, "foto_selfie_url": data_url_foto, "tipe_absen": menu_kru}
                    http_session.post(API_URL, json=payload)
                    if is_valid: st.success(f"Absensi valid. Selisih: {jarak:.1f} m")
                    else: st.warning(f"Posisi di luar batas ({jarak:.1f} m).")
                    time.sleep(2); st.rerun()

    elif menu_kru == "To-Do List Divisi":
        st.subheader("📋 Lembar Kerja & Verifikasi Tugas")
        df_all_tasks = fetch_data("tasks")

        if not df_all_tasks.empty and 'crew_name' in df_all_tasks.columns:
            my_tasks = df_all_tasks[(df_all_tasks['crew_name'] == st.session_state.user_name) & (df_all_tasks['event_id'] == active_ev_id)]

            tasks_todo = my_tasks[my_tasks['status'].isin(["Belum Selesai", "Ditolak"])]
            tasks_pending = my_tasks[my_tasks['status'] == "Menunggu Verifikasi"]
            tasks_approved = my_tasks[my_tasks['status'] == "Disetujui"]

            st.markdown("### 🔴 Perlu Dikerjakan / Ditolak")
            if tasks_todo.empty: st.info("Tidak ada tugas baru.")
            else:
                for idx, row in tasks_todo.iterrows():
                    with st.container():
                        st.markdown(f"**Tugas:** {row['task_name']}")
                        if row['status'] == "Ditolak": st.error(f"⚠️ DITOLAK TL. Alasan: {row.get('tl_comment', '-')}")

                        if pd.notna(row.get('reference_photo', '')) and "base64" in str(row['reference_photo']):
                            try:
                                b64_str = str(row['reference_photo']).split("base64,")[1]
                                st.image(Image.open(io.BytesIO(base64.b64decode(b64_str))), caption="📸 Arahan Visual dari TL")
                            except: pass

                        cam_report = st.camera_input(f"Ambil Bukti Selesai", key=f"cam_report_{row['task_id']}")
                        if st.button(f"Kirim Laporan: {row['task_name']}", key=f"btn_done_{row['task_id']}", use_container_width=True):
                            if not cam_report: st.error("Harap foto bukti kerjaan.")
                            else:
                                with st.spinner("Mengunggah bukti ke TL..."):
                                    raw_img = Image.open(io.BytesIO(cam_report.getvalue()))
                                    raw_img.thumbnail((300, 300))
                                    buffered_io = io.BytesIO()
                                    raw_img.save(buffered_io, format="JPEG", quality=50)
                                    encoded_report = f"data:image/jpeg;base64,{base64.b64encode(buffered_io.getvalue()).decode('utf-8')}"

                                    payload_update = {"action": "update_task", "task_id": row['task_id'], "status": "Menunggu Verifikasi", "photo_report": encoded_report}
                                    http_session.post(API_URL, json=payload_update)
                                    st.success("Tugas dikirim ke TL untuk diverifikasi."); time.sleep(1); st.rerun()
                        st.markdown("---")

            st.markdown("### 🟡 Menunggu Verifikasi TL")
            if tasks_pending.empty: st.info("Tidak ada laporan yang sedang diperiksa.")
            else:
                for idx, row in tasks_pending.iterrows():
                    st.warning(f"⏳ **{row['task_name']}** - Sedang ditinjau oleh Team Leader.")

            st.markdown("### 🟢 Selesai & Disetujui")
            if tasks_approved.empty: st.info("Belum ada tugas yang disetujui.")
            else:
                for idx, row in tasks_approved.iterrows():
                    st.success(f"✅ **{row['task_name']}**")
                    if pd.notna(row.get('tl_comment', '')) and row['tl_comment']:
                        st.caption(f"Komentar TL: {row['tl_comment']}")
        else: st.info("Belum ada alokasi tugas untuk Anda.")

        st.markdown("---")
        st.subheader("👥 Papan Status Rekan Divisi")
        if not df_all_tasks.empty and 'division' in df_all_tasks.columns:
            div_tasks = df_all_tasks[(df_all_tasks['event_id'] == active_ev_id) & (df_all_tasks['division'] == my_division)]
            if not div_tasks.empty: st.dataframe(div_tasks[['crew_name', 'task_name', 'status']], use_container_width=True)
            else: st.info("Belum ada alokasi pekerjaan untuk divisi Anda.")
        else: st.info("Sistem belum mendeteksi riwayat pekerjaan dari pusat.")
